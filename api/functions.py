import re
from oauth2_provider.models import AccessToken
from django.utils import timezone
from PyPDF2 import PdfReader
from transformers import AutoModelForQuestionAnswering, AutoTokenizer, T5Tokenizer, T5ForConditionalGeneration, AutoModelForCausalLM
from googletrans import Translator
from sentence_transformers import SentenceTransformer, util
from docx import Document
import zipfile
import os
from openpyxl import load_workbook
import win32com.client as win32
import xlsxwriter
from docx.shared import Inches
import pythoncom
import random

def translate(texto, src="es", dest="en"): 
  try:
    traductor = Translator()
    traduccion = traductor.translate(texto, src=src, dest=dest)
    return traduccion.text
  except TypeError:
    return None

def process_pdf_and_generate_qa(file_path):
  # Variable que contendrá las preguntas y respuestas
  preguntas_respuestas = []
  
  # Modelos
  question_model_name = "mrm8488/t5-base-finetuned-question-generation-ap"
  qa_model_name = "deepset/roberta-base-squad2"
  
  question_tokenizer = T5Tokenizer.from_pretrained(question_model_name)
  question_model = T5ForConditionalGeneration.from_pretrained(question_model_name)
  
  qa_tokenizer = AutoTokenizer.from_pretrained(qa_model_name)
  qa_model = AutoModelForQuestionAnswering.from_pretrained(qa_model_name)
  
  # Cargar el archivo PDF
  reader = PdfReader(file_path)
  preguntas_traducidas = []
  
  for page in reader.pages:
    # Extraer texto de la página
    texto = page.extract_text()
    
    # Limpiar el texto y traducir al inglés
    texto_limpio = translate(re.sub(r'\s+', ' ', texto.strip()))
    if (texto_limpio is not None):
      # Preparar el texto para generar preguntas
      input_text = f"context: {texto_limpio}"
      inputs = question_tokenizer.encode(input_text, return_tensors="pt", truncation=True)
      
      # Generar preguntas
      outputs = question_model.generate(
        inputs,
        max_length=64,
        num_return_sequences=1,
        num_beams=10,
        early_stopping=True
      )
      
      for output in outputs:
        # Pregunta traducida al español
        pregunta_a_traducir = question_tokenizer.decode(output, skip_special_tokens=True)
        if (pregunta_a_traducir not in preguntas_traducidas):
          preguntas_traducidas.append(pregunta_a_traducir)
          pregunta = translate(pregunta_a_traducir, src="en", dest="es")
          if (pregunta is not None):
            # Usar modelo QA para generar respuesta
            qa_inputs = qa_tokenizer(
              translate(pregunta, src="es", dest="en"),  # Pregunta en inglés
              texto_limpio,  # Contexto en inglés
              return_tensors="pt",
              truncation=True
            )
            qa_outputs = qa_model(**qa_inputs)
            
            # Extraer respuesta del modelo
            respuesta_start = qa_outputs.start_logits.argmax()
            respuesta_end = qa_outputs.end_logits.argmax()
            respuesta = qa_tokenizer.convert_tokens_to_string(
              qa_tokenizer.convert_ids_to_tokens(qa_inputs["input_ids"][0][respuesta_start:respuesta_end + 1])
            )
            
            # Traducir respuesta al español
            respuesta_correcta = translate(respuesta, src="en", dest="es")
            if (respuesta_correcta is not None):
              # Guardar la pregunta y la respuesta
              preguntas_respuestas.append({
                "pregunta": pregunta,
                "respuesta_correcta": respuesta_correcta
              })
  return preguntas_respuestas
  
def get_user_from_token(token: str):
  try:
    # Get the AccessToken object
    access_token = AccessToken.objects.get(token=token)

    # Check if the token has expired
    if access_token.expires < timezone.now():
        return None

    # Return the associated user
    return access_token.user
  except AccessToken.DoesNotExist:
    return None

# Evaluación de la respuesta del usuario
model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
def evaluate_user_response(respuesta_correcta, respuesta_usuario):
  # Generar embeddings
  embedding_correcta = model.encode(respuesta_correcta, convert_to_tensor=True)
  embedding_usuario = model.encode(respuesta_usuario, convert_to_tensor=True)
  
  # Calcular similitud
  similaridad = util.cos_sim(embedding_correcta, embedding_usuario).item()

  # Escalar a un rango de 1-10
  puntaje = round(similaridad * 10, 1)

  return puntaje

# Cargar el modelo para generar retroalimentación
feedback_model_name = "EleutherAI/gpt-neo-1.3B"  # Modelo de generación de texto
feedback_tokenizer = AutoTokenizer.from_pretrained(feedback_model_name)
feedback_model = AutoModelForCausalLM.from_pretrained(feedback_model_name)

def generate_feedback(responses, correct_answers, scores, user_name="Usuario"):
  # Crear un resumen del desempeño del usuario
  performance_summary = []
  for i, (user_response, correct_answer, score) in enumerate(zip(responses, correct_answers, scores)):
    performance_summary.append(
      f"Pregunta {i + 1}: Respuesta del usuario: '{user_response}'. Respuesta correcta: '{correct_answer}'. Puntaje: {score}/10."
    )
  
  # Crear el prompt para el modelo
  prompt = (
    f"El estudiante {user_name} ha completado una evaluación. Aquí están los detalles:\n\n"
    f"{' '.join(performance_summary)}\n\n"
    "Con base en este desempeño, proporcione recomendaciones personalizadas para mejorar el conocimiento del usuario en los temas evaluados."
  )
  
  # Tokenizar e inferir
  inputs = feedback_tokenizer.encode(prompt, return_tensors="pt", max_length=512, truncation=True)
  outputs = feedback_model.generate(inputs, max_length=300, num_return_sequences=1, no_repeat_ngram_size=2)
  
  # Decodificar la recomendación
  feedback = feedback_tokenizer.decode(outputs[0], skip_special_tokens=True)
  return feedback

# ---------------------------------------------- #
# Función para reemplazar texto en un documento
def replace_text_in_word(file_path, new_file_path, replacements):
  print(f"replace_text_in_word({file_path}, {new_file_path})")
  # Abrir el documento existente
  doc = Document(file_path)
  
  # Recorrer todos los párrafos del documento
  for paragraph in doc.paragraphs:
    for old_text, new_text in replacements.items():
      if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(old_text, new_text)
  
  # Guardar el documento modificado
  doc.save(new_file_path)

# Función para extraer el contenido del archivo .docx
def extract_docx(docx_path, extract_to):
  with zipfile.ZipFile(docx_path, 'r') as docx:
    docx.extractall(extract_to)

# Función para modificar los datos en el archivo Excel incrustado
def modify_excel_data(excel_path, chart_index):
  # Cargar el archivo Excel incrustado
  wb = load_workbook(excel_path)
  sheet = wb.active

  # Lógica de modificación de datos (esto puede personalizarse según el gráfico)
  if chart_index == 0:
    # Cambiar los datos del primer gráfico
    sheet['A2'] = 15
    sheet['B2'] = 25
    sheet['C2'] = 35
  elif chart_index == 1:
    # Cambiar los datos del segundo gráfico
    sheet['A2'] = 50
    sheet['B2'] = 60
    sheet['C2'] = 70
  else:
    # Modificar según el gráfico específico
    sheet['A2'] = 100
    sheet['B2'] = 120
    sheet['C2'] = 140

  # Guardar los cambios en el archivo Excel
  wb.save(excel_path)

# Función para volver a empaquetar el archivo .docx con el Excel modificado
def update_docx_with_modified_excel(docx_path, extracted_folder, new_file_path):
  with zipfile.ZipFile(new_file_path, 'w', zipfile.ZIP_DEFLATED) as docx:
    for root, dirs, files in os.walk(extracted_folder):
      for file in files:
        file_path = os.path.join(root, file)
        docx.write(file_path, os.path.relpath(file_path, extracted_folder))

# ----------------------------------------------------------------------------------------- #
def modify_word_document(file_path, new_file_path, replacements):
  print("modify_word_document")
  # Paso 1: Extraer el contenido del archivo .docx
  # extracted_folder = 'extracted'
  # extract_docx(file_path, extracted_folder)

  # Paso 2: Buscar y modificar los gráficos (archivos Excel incrustados)
  # excel_file_paths = []
  # for root, dirs, files in os.walk(extracted_folder):
  #   for file in files:
  #     if file.endswith('.xlsx'):
  #       excel_file_paths.append(os.path.join(root, file))

  # Modificar los datos en los archivos Excel (gráficos)
  # for i, excel_file_path in enumerate(excel_file_paths):
  #   modify_excel_data(excel_file_path, i)

  # Paso 3: Volver a empaquetar el archivo .docx con los cambios en los gráficos
  # update_docx_with_modified_excel(file_path, extracted_folder, new_file_path)
  
  # Paso 4: Reemplazar texto en el documento
  replace_text_in_word(file_path, new_file_path, replacements)
  
  # Paso 5: Borrar directorio de extracted
  # shutil.rmtree("extracted")
  
  # Paso 6: Actualizar gráficos
  update_graphics_in_word(new_file_path, [
    [
      ['Categoría', 'Matemática', 'Seguridad de Software', 'Base de datos'],
      ['04/11/2024 - 10/11/2024', 1, 2, 3],
      ['11/11/2024 - 17/11/2024', 2, 0, 4],
      ['18/11/2024 - 24/11/2024', 5, 2, 3],
      ['25/11/2024 - 01/12/2024', 9, 1, 2],
    ],
    [
      ['Categoría', 'Matemática', 'Seguridad de Software', 'Base de datos'],
      ['04/11/2024 - 10/11/2024', 22, 6, 10],
      ['11/11/2024 - 17/11/2024', 12, 10, 7],
      ['18/11/2024 - 24/11/2024', 9, 12, 15],
      ['25/11/2024 - 01/12/2024', 16, 24, 10],
    ],
    [
      ['Categoría', 'Matemática', 'Seguridad de Software', 'Base de datos'],
      ['04/11/2024 - 10/11/2024', 1, 2, 3],
      ['11/11/2024 - 17/11/2024', 2, 0, 4],
      ['18/11/2024 - 24/11/2024', 5, 2, 3],
      ['25/11/2024 - 01/12/2024', 9, 1, 2],
    ],
    [
      ['Categoría', 'Matemática', 'Seguridad de Software', 'Base de datos'],
      ['04/11/2024 - 10/11/2024', 20, 4, 9],
      ['11/11/2024 - 17/11/2024', 10, 9, 7],
      ['18/11/2024 - 24/11/2024', 9, 10, 12],
      ['25/11/2024 - 01/12/2024', 15, 19, 6],
    ],
    [
      ['Categoría', 'Matemática', 'Seguridad de Software', 'Base de datos'],
      ['04/11/2024 - 10/11/2024', 2, 2, 1],
      ['11/11/2024 - 17/11/2024', 2, 1, 0],
      ['18/11/2024 - 24/11/2024', 0, 2, 3],
      ['25/11/2024 - 01/12/2024', 1, 5, 4],
    ],
    [
      ['Categoría', 'Matemática', 'Seguridad de Software', 'Base de datos'],
      ['Porcentaje de errores', 9, 23, 23],
    ],
    [
      ['Categoría', 'Matemática', 'Seguridad de Software', 'Base de datos'],
      ['04/11/2024 - 10/11/2024', 8.5, 6.3, 8.5],
      ['11/11/2024 - 17/11/2024', 8.2, 8.5, 10],
      ['18/11/2024 - 24/11/2024', 10, 8.2, 8],
      ['25/11/2024 - 01/12/2024', 9.8, 8, 6],
    ],
  ])
  
def update_graphics_in_word(docx_path, data):
  print("update_graphics_in_word")
  replacements = {}
  print("Flag 1")
  
  # Generar el gráfico e imagen para cada conjunto de datos
  for i, dataset in enumerate(data, start=1):
    print("Flag 2")
    # Generar el gráfico y guardarlo como imagen
    image_path = generate_chart_and_save_image(dataset, i)
    print("Flag 3")
    # Asociar el marcador de texto {{graficoX}} con la imagen generada
    replacements[f'var_grafico_{i}'] = image_path
    print("Flag 4")
  # Reemplazar las imágenes en el documento de Word
  print("Flag 5")
  replace_image_in_word(docx_path, docx_path, replacements)

def generate_chart_and_save_image(data, chart_number):
  print("generate_chart_and_save_image")
  # Inicializar COM
  pythoncom.CoInitialize()

  # Crear un archivo de Excel temporal con una ruta absoluta
  file_name = f"grafico_columnas_agrupadas_{chart_number}.xlsx"
  file_path = os.path.abspath(file_name)  # Obtener la ruta absoluta del archivo
  workbook = xlsxwriter.Workbook(file_path)
  worksheet = workbook.add_worksheet()

  # Escribir los datos en el archivo
  for row_num, row_data in enumerate(data):
    worksheet.write_row(row_num, 0, row_data)

  # Crear el gráfico de columnas agrupadas
  chart = workbook.add_chart({'type': 'column' if chart_number != 7 else 'line'})
  
  # Aplicar un estilo moderno al gráfico
  chart.ChartStyle = 250

  # Agregar las series al gráfico
  var_len = (len(data) - 1) if chart_number != 6 else len(data)
  
  for col_num in range(1, len(data[0])):  # Recorre las columnas desde la segunda
    # Escribe las categorías en la primera columna (columna A)
    categories_start_row = 1  # Start from row 2 (index 1) because row 1 has headers
    categories_end_row = len(data) - 1  # Last row with data
    categories_range = f'=Sheet1!$A${categories_start_row + 1}:$A${categories_end_row + 1}'

    # Escribe los valores en la columna actual
    values_start_row = categories_start_row
    values_end_row = categories_end_row
    values_range = f'=Sheet1!${chr(65 + col_num)}${values_start_row + 1}:${chr(65 + col_num)}${values_end_row + 1}'

    # Add series to the chart
    chart.add_series({
      'name': f'=Sheet1!${chr(65 + col_num)}$1',  # Header row for the column
      'categories': categories_range,
      'values': values_range,
    })

  # Configurar el gráfico
  chart.set_title({'name': 'Gráfico de Columnas Agrupadas'})
  chart.set_x_axis({'name': 'Mes'})
  chart.set_y_axis({'name': 'Valor'})

  # Insertar el gráfico en la hoja de trabajo
  worksheet.insert_chart('E2', chart)

  # Guardar el archivo de Excel
  workbook.close()

  # Convertir el gráfico en Excel a una imagen
  excel = win32.Dispatch('Excel.Application')
  excel.Visible = False
  workbook = excel.Workbooks.Open(file_path)
  sheet = workbook.Sheets('Sheet1')
  chart = sheet.ChartObjects(1).Chart  # Seleccionar el gráfico

  # Desactivar Título del gráfico
  chart.HasTitle = False

  # Desactivar Títulos de los ejes
  chart.Axes(1).HasTitle = False  # Eje X
  chart.Axes(2).HasTitle = False  # Eje Y

 # Desactivar etiquetas de datos para cada serie
  for serie in chart.SeriesCollection():
    serie.HasDataLabels = False  # Desactivar etiquetas de datos
    if serie.HasErrorBars:
      serie.ErrorBars().Delete()

  # Activar Tabla de datos
  chart.HasDataTable = True

  # Activar Líneas de cuadrícula
  chart.Axes(1).HasMajorGridlines = True  # Líneas de cuadrícula en el eje X
  chart.Axes(2).HasMajorGridlines = True  # Líneas de cuadrícula en el eje Y

  # Activar Leyenda
  chart.HasLegend = False

  # Desactivar Línea de tendencia (si existe)
  try:
    chart.SeriesCollection(1).Trendlines.Delete()
  except:
    pass

  # Exportar el gráfico como imagen PNG
  image_path = f"grafico_columnas_agrupadas_{chart_number}.png"
  image_path_abs = os.path.abspath(image_path)  # Obtener la ruta absoluta de la imagen
  chart.Export(image_path_abs)

  # Cerrar Excel
  workbook.Close(False)
  excel.Quit()

  # Desinicializar COM
  pythoncom.CoUninitialize()

  return image_path_abs

def replace_image_in_word(docx_path, new_file_path, replacements):
  print("replace_image_in_word")
  # Abrir el documento existente
  doc = Document(docx_path)
  
  # Recorrer todos los párrafos del documento
  for paragraph in doc.paragraphs:
    for old_text, image_path in replacements.items():
      print(f"if {old_text} in {paragraph.text}:")
      if old_text in paragraph.text:
        paragraph.clear()
        paragraph.add_run().add_picture(image_path, width=Inches(6), height=Inches(3.5))
        print(f"os.remove: {image_path.replace('png', 'xlsx')}")
        print(f"os.remove: {image_path}")
        os.remove(image_path.replace("png", "xlsx"))
        os.remove(image_path)
        
  # Guardar el documento modificado
  doc.save(new_file_path)
# ----------------------------------------------------------------------------------------- #

# Función para generar código de 5 digitos
def generate_otp():
  return str(random.randint(100000, 999999))  # Genera un código de 6 dígitos