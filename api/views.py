from io import BytesIO
import os
import tempfile
import comtypes
from django.http import FileResponse, HttpResponseBadRequest
from django.shortcuts import get_object_or_404
from django.forms import ValidationError
from docx import Document
import comtypes.client
from rest_framework import viewsets, status
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.response import Response
from rest_framework.decorators import action
from rest_framework.views import APIView
from django.contrib.auth.hashers import make_password, check_password
from oauth2_provider.models import AccessToken
from django.utils import timezone
from oauth2_provider.settings import oauth2_settings
from oauthlib.common import generate_token
from datetime import time, datetime, timedelta
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from django.core.mail import send_mail
from PIL import Image
from .models import (
  OTPCode,
  Rol, 
  Usuario, 
  Periodo, 
  Materia,
  MateriaPeriodo,
  DocumentoPDF,
  Evaluacion,
  Pregunta,
  Recomendacion,
  Historial, 
  Notificacion
)
from .serializers import (
  RolSerializer, 
  UsuarioSerializer, 
  PeriodoSerializer, 
  MateriaSerializer,
  MateriaPeriodoSerializer,
  DocumentoPDFSerializer,
  EvaluacionSerializer,
  PreguntaSerializer,
  RecomendacionSerializer,
  HistorialSerializer, 
  NotificacionSerializer
)
from .functions import (
  generate_feedback,
  modify_word_document,
  process_pdf_and_generate_qa,
  evaluate_user_response,
  generate_otp
)

class ValidateTokenAPIView(APIView):
  def get(self, request):
    token_str = request.headers.get("Authorization", "").split("Bearer ")[-1]
    try:
      token = AccessToken.objects.get(token=token_str)
      if token.expires < timezone.now():
        return Response({"error": "El token ha expirado"}, status=status.HTTP_401_UNAUTHORIZED)
      return Response({"message": "Token válido"}, status=status.HTTP_200_OK)
    except AccessToken.DoesNotExist:
      return Response({"error": "Token inválido"}, status=status.HTTP_401_UNAUTHORIZED)

class RolViewSet(viewsets.ModelViewSet):
  queryset = Rol.objects.all()
  serializer_class = RolSerializer
  permission_classes=[IsAuthenticated]

class PeriodoViewSet(viewsets.ModelViewSet):
  queryset = Periodo.objects.all()
  serializer_class = PeriodoSerializer

  def get_permissions(self):
    if self.action == 'list':
        self.permission_classes = [AllowAny]
    else:
        self.permission_classes = [IsAuthenticated]
    return super().get_permissions()

class MateriaViewSet(viewsets.ModelViewSet):
  queryset = Materia.objects.all()
  serializer_class = MateriaSerializer
  permission_classes=[IsAuthenticated]

class MateriaPeriodoViewSet(viewsets.ModelViewSet):
  queryset = MateriaPeriodo.objects.all()
  serializer_class = MateriaPeriodoSerializer
  permission_classes=[IsAuthenticated]
  
  def get_queryset(self):
    queryset = MateriaPeriodo.objects.all()
    # Obtener los parámetros de consulta
    materia = self.request.query_params.get('materia')
    periodo = self.request.query_params.get('periodo')
    # Filtrar por materia si se proporciona
    if materia is not None:
      queryset = queryset.filter(materia=materia)
    # Filtrar por periodo si se proporciona
    if periodo is not None:
      queryset = queryset.filter(periodo=periodo)
    return queryset

class UsuarioViewSet(viewsets.ModelViewSet):
  queryset = Usuario.objects.all()
  serializer_class = UsuarioSerializer
  permission_classes=[IsAuthenticated]
  
  def get_permissions(self):
    if self.action in ['create', 'login', 'update', 'partial_update']:
      return [AllowAny()]
    return [permission() for permission in self.permission_classes]
  
  def perform_create(self, serializer):
    password = serializer.validated_data['password']
    hashed_password = make_password(password)
    serializer.save(password=hashed_password)
  
  def update(self, request, *args, **kwargs):
    usuario = self.get_object()
    data = request.data.copy()  # Copia los datos para evitar modificar el original
    # Procedimientos antes de actualizar
    if "password" in data:
      hashed_password = make_password(data["password"])
      usuario.password = hashed_password  # Guardar directamente en el modelo
      usuario.save(update_fields=['password'])
      data.pop("password")  # Remover del diccionario para que no lo sobrescriba el serializer
    if "email" in data:
      existing_user = Usuario.objects.filter(email=data["email"]).exclude(id=usuario.id).first()
      if existing_user:
        return Response({"error": "El email ya está en uso"}, status=status.HTTP_400_BAD_REQUEST)

    # Continua con la actualización sin el campo "password"
    return super().update(request, *args, **kwargs)

  def partial_update(self, request, *args, **kwargs):
    print("partial_update")
    usuario = self.get_object()
    data = request.data.copy()  # Copia los datos para evitar modificar el original
    # Validar email
    if "email" in data:
      existing_user = Usuario.objects.filter(email=data["email"]).exclude(id=usuario.id).first()
      if existing_user:
        return Response({"error": "El email ya está en uso"}, status=status.HTTP_400_BAD_REQUEST)
    # Encriptar la contraseña antes de actualizar
    if "password" in data:
      hashed_password = make_password(data["password"])
      usuario.password = hashed_password  # Guardar la contraseña encriptada
      usuario.save(update_fields=['password'])  # Guardar solo el campo de contraseña
      request.data.pop("password")  # Eliminar la contraseña del diccionario para que no sea sobrescrita
    # Continúa con la actualización parcial sin el campo "password"
    return super().partial_update(request, *args, **kwargs)

  @action(detail=False, methods=['get'], url_path="email", permission_classes=[AllowAny])
  def email(self, request):
    email = request.GET.get("email")
    if not email:
      return Response({"error": "Email es requerido"}, status=status.HTTP_400_BAD_REQUEST)
    
    user = get_object_or_404(Usuario, email=email)
    return Response({"id": user.id}, status=status.HTTP_200_OK)
  
  @action(detail=False, methods=['post'], url_path="login", permission_classes=[AllowAny])
  def login(self, request):
    id = request.data.get("id")
    email = request.data.get("email")
    password = request.data.get("password")
    try:
      if (id):
        user = Usuario.objects.get(id=id)
        token = AccessToken.objects.create(
          user=user,
          scope='read write',
          expires=timezone.now() + timedelta(seconds=oauth2_settings.ACCESS_TOKEN_EXPIRE_SECONDS),
          token=generate_token(),
          application=None
        )
        return Response({
          "access_token": token.token,
          "expires_in": oauth2_settings.ACCESS_TOKEN_EXPIRE_SECONDS,
          "token_type": "Bearer",
        }, status=status.HTTP_200_OK)
      else:
        user = Usuario.objects.get(email=email)
        if check_password(password, user.password):
          token = AccessToken.objects.create(
            user=user,
            scope='read write',
            expires=timezone.now() + timedelta(seconds=oauth2_settings.ACCESS_TOKEN_EXPIRE_SECONDS),
            token=generate_token(),
            application=None
          )
          return Response({
            "access_token": token.token,
            "expires_in": oauth2_settings.ACCESS_TOKEN_EXPIRE_SECONDS,
            "token_type": "Bearer",
          }, status=status.HTTP_200_OK)
        else:
          return Response({"error": "Credenciales inválidas"}, status=status.HTTP_400_BAD_REQUEST)
    except Usuario.DoesNotExist:
        return Response({"error": "Usuario no encontrado"}, status=status.HTTP_404_NOT_FOUND)
  
  @action(detail=False, methods=['post'], url_path="report")
  def report(self, request):
    id = request.data.get("id")
    try:
      if (id):
        usuario = Usuario.objects.get(id=id)
        input_file = "template.docx"  # Archivo original
        output_file = f"Reporte_{usuario.nombres}_{usuario.apellidos}.docx"  # Archivo modificado
        
        # Obtener los documentos PDF del usuario
        documentos_pdf = DocumentoPDF.objects.filter(usuario__id=id)
        
        # Obtener todas las evaluaciones correspondientes a esos documentos PDF
        evaluaciones = Evaluacion.objects.filter(documento_pdf__in=documentos_pdf)
        
        # Obtener todas las evaluaciones realizadas
        evaluaciones_realizadas = Evaluacion.objects.filter(duracion="00:00:00")

        # Obtener todas las preguntas relacionadas con esas evaluaciones
        preguntas = Pregunta.objects.filter(evaluacion__in=evaluaciones)
        
        # Obtener todas las preguntas con respuestas correctas
        preguntas_correctas = Pregunta.objects.filter(evaluacion__in=evaluaciones, puntaje__gte=5)
        
        # Obtener todas las preguntas con respuestas correctas
        preguntas_incorrectas = Pregunta.objects.filter(evaluacion__in=evaluaciones, puntaje__lt=5)

        # Datos a reemplazar
        replacements = {
          "var_nombre": f"{usuario.nombres} {usuario.apellidos}",
          "var_cedula": f"{str(usuario.cedula)}",
          "var_periodo": f"{usuario.periodo.nombre}",
          "var_pdfs": str(len(documentos_pdf)), 
          "var_preguntas": str(len(preguntas)), 
          "var_evaluaciones_realizadas": str(len(evaluaciones_realizadas)), 
          "var_promedio_respuestas_correctas": str((len(preguntas_correctas) / len(preguntas)) * 100),
          "var_analisis_preguntas_fallidas": str((len(preguntas_incorrectas) / len(preguntas)) * 100),
        }

        # Llamar a la función
        modify_word_document(os.path.abspath(input_file), os.path.abspath(output_file), replacements)
        return Response({}, status=status.HTTP_201_CREATED)
      else:
        return Response({"error": "Evaluación no ingresada"}, status=status.HTTP_400_BAD_REQUEST)
    except Evaluacion.DoesNotExist:
        return Response({"error": "Usuario no encontrado"}, status=status.HTTP_404_NOT_FOUND)

class DocumentoPDFViewSet(viewsets.ModelViewSet):
  queryset = DocumentoPDF.objects.all()
  serializer_class = DocumentoPDFSerializer
  permission_classes = [IsAuthenticated]

  def get_queryset(self):
    queryset = DocumentoPDF.objects.all()
    usuario = self.request.query_params.get('usuario')
    materia_periodo = self.request.query_params.get('materia-periodo')
    orden = self.request.query_params.get('orden')

    if usuario and materia_periodo:
      queryset = queryset.filter(usuario=usuario, materia_periodo=materia_periodo)
    else:
      if usuario is not None:
        queryset = queryset.filter(usuario=usuario)
      if materia_periodo is not None:
        queryset = queryset.filter(materia_periodo=materia_periodo)

    if orden:
      if orden == 'fecha-carga':
        queryset = queryset.order_by('-fecha')
      elif orden == 'nombre':
        queryset = queryset.order_by('archivo')
      else:
        raise ValidationError(f"El valor de 'orden' '{orden}' no es válido. Debe ser 'fecha-carga', o 'nombre'.")
    return queryset

  def perform_create(self, serializer):
    instance = serializer.save()
    self.process_evaluation_and_notifications(instance, serializer)
  
  def perform_update(self, serializer):
    instance = serializer.save()
    self.process_evaluation_and_notifications(instance, serializer)
    
  def process_evaluation_and_notifications(self, instance, serializer, update=False):
    fecha_evaluacion = serializer.initial_data.get('fecha_evaluacion', None)
    file_path = instance.archivo.path

    # Si es actualización, obtenemos la evaluación existente o creamos una nueva
    if update:
      evaluacion, created = Evaluacion.objects.get_or_create(documento_pdf=instance)
      evaluacion.documento_pdf = instance
      evaluacion.titulo = os.path.splitext(os.path.basename(file_path))[0]
      evaluacion.fecha_evaluacion = fecha_evaluacion
      evaluacion.duracion = time(hour=1, minute=0)  # Se mantiene 60 min por defecto
      evaluacion.save()
      # Creamos un registro en historial
      Historial.objects.create(
        usuario=evaluacion.documento_pdf.usuario,
        descripcion=f'Archivo actualizado en la materia {evaluacion.documento_pdf.materia_periodo.materia.nombre}',
        fecha=datetime.now()
      )

      # Eliminar preguntas antiguas antes de regenerar
      evaluacion.preguntas.all().delete()
    else:
      evaluacion, created = Evaluacion.objects.get_or_create(
        documento_pdf=instance,
        defaults={
          "documento_pdf": instance,
          "titulo": os.path.splitext(os.path.basename(file_path))[0],
          "fecha_evaluacion": fecha_evaluacion,
          "duracion": time(hour=1, minute=0)
        }
      )
      if not created:
        # Si ya existe, solo actualiza la fecha
        evaluacion.documento_pdf = instance
        evaluacion.titulo = os.path.splitext(os.path.basename(file_path))[0]
        evaluacion.fecha_evaluacion = fecha_evaluacion
        evaluacion.duracion = time(hour=1, minute=0)  # Se mantiene 60 min por defecto
        evaluacion.save()
      # Creamos un registro en historial
      Historial.objects.create(
        usuario=evaluacion.documento_pdf.usuario,
        descripcion=f'Archivo subido en la materia {evaluacion.documento_pdf.materia_periodo.materia.nombre}',
        fecha=datetime.now()
      )

    # Eliminar las preguntas actuales
    Pregunta.objects.filter(evaluacion=evaluacion).delete()

    # Procesar el PDF y generar preguntas
    questions_answers = process_pdf_and_generate_qa(file_path)
    for qa in questions_answers:
      Pregunta.objects.create(
        evaluacion=evaluacion,
        pregunta=qa["pregunta"],
        respuesta_correcta=qa["respuesta_correcta"],
        respuesta="",
        puntaje=0.0
      )

    # Manejo de notificaciones
    usuario = get_object_or_404(Usuario, id=evaluacion.documento_pdf.usuario.id)
    if fecha_evaluacion:
      try:
        fecha_evaluacion_datetime = datetime.strptime(fecha_evaluacion, '%Y-%m-%dT%H:%M')
        fecha_actual = datetime.now()
        diferencia_tiempo = fecha_evaluacion_datetime - fecha_actual
        intervalo = diferencia_tiempo / 9

        # Eliminar notificaciones previas si se está actualizando
        if update:
          Notificacion.objects.filter(usuario=usuario, evaluacion=evaluacion.id).delete()

        # Crear nuevas notificaciones
        for i in range(10):
          fecha_notificacion = fecha_actual + intervalo * i
          Notificacion.objects.create(
            usuario=usuario,
            evaluacion=evaluacion,
            titulo=evaluacion.documento_pdf.materia_periodo.materia.nombre,
            descripcion=f"Recordatorio de estudio - {evaluacion.documento_pdf.archivo.name.split('/').pop()}",
            fecha_creacion=datetime.now(),
            fecha_aparicion=fecha_notificacion,
            url="#",
            visto=False
          )

      except ValueError as e:
          print(f"Error en la fecha de evaluación: {e}")

  def destroy(self, request, *args, **kwargs):
    instance = self.get_object()
    file_path = instance.archivo.path
    if os.path.isfile(file_path):
      os.remove(file_path)
    self.perform_destroy(instance)
    return Response(status=status.HTTP_204_NO_CONTENT)

class EvaluacionViewSet(viewsets.ModelViewSet):
  queryset = Evaluacion.objects.all()
  serializer_class = EvaluacionSerializer
  permission_classes = [IsAuthenticated]
  
  @action(detail=False, methods=['post'], url_path="check")
  def check(self, request):
    id = request.data.get("id")
    try:
      if (id):
        evaluacion = Evaluacion.objects.get(id=id)
        preguntas = Pregunta.objects.filter(evaluacion=id)
        responses = [pregunta.respuesta for pregunta in preguntas]
        correct_answers = [pregunta.respuesta_correcta for pregunta in preguntas]
        scores = [pregunta.puntaje for pregunta in preguntas]
        feedback = generate_feedback(responses=responses, correct_answers=correct_answers, scores=scores)
        print(feedback)
      else:
        return Response({"error": "Evaluación no ingresada"}, status=status.HTTP_400_BAD_REQUEST)
    except Evaluacion.DoesNotExist:
        return Response({"error": "Usuario no encontrado"}, status=status.HTTP_404_NOT_FOUND)

  def get_queryset(self):
    queryset = Evaluacion.objects.all()

    # Obtener los parámetros de consulta
    documento_pdf = self.request.query_params.get('documento_pdf')
    usuario = self.request.query_params.get('usuario')

    # Filtrar por documento_pdf si se proporciona
    if documento_pdf is not None:
      queryset = queryset.filter(documento_pdf=documento_pdf)

    # Filtrar por usuario si se proporciona
    if usuario is not None:
      queryset = queryset.filter(documento_pdf__usuario=usuario)

    return queryset

class PreguntaViewSet(viewsets.ModelViewSet):
  queryset = Pregunta.objects.all()
  serializer_class = PreguntaSerializer
  permission_classes = [IsAuthenticated]

  def get_queryset(self):
    queryset = Pregunta.objects.all()
    evaluacion = self.request.query_params.get('evaluacion')
    usuario = self.request.query_params.get('usuario')
    if evaluacion is not None:
      queryset = queryset.filter(evaluacion=evaluacion)
    if usuario is not None:
      queryset = queryset.filter(evaluacion__documento_pdf__usuario=usuario)
    return queryset

  def update(self, request, *args, **kwargs):
    from google import genai
    client = genai.Client(api_key="AIzaSyCT0e2R7V-Pt3pdUy2Oy6Htm7uQmglN9GQ")
    # Obtener la instancia antes de la actualización
    instance = self.get_object()
    
    # Respuestas
    respuesta_usuario = request.data.get("respuesta", None)
    puntaje = request.data.get("puntaje", None)
    
    Recomendacion.objects.filter(pregunta=instance.id).exclude(id=instance.id).delete()

    # Evaluar las respuestas
    if (puntaje == 0):
      contents = f"""
        Dame una recomendación con respecto a: El estudiante seleccionó la respuesta '{respuesta_usuario}' y la pregunta MCQ es '{instance.pregunta}'.
      """
      response = client.models.generate_content(
          model="gemini-2.0-flash",
          contents=contents,
      )
      Recomendacion.objects.create(
        materia_periodo=instance.evaluacion.documento_pdf.materia_periodo,
        usuario=instance.evaluacion.documento_pdf.usuario,
        evaluacion=instance.evaluacion,
        pregunta=instance,
        descripcion=response.text
      )
    
    instance.puntaje = puntaje
    
    # Guardar los cambios
    instance.save()

    # Llamar al método original para realizar la actualización
    response = super().update(request, *args, **kwargs)

    # Retornar la respuesta
    return response

class RecomendacionViewSet(viewsets.ModelViewSet):
  queryset = Recomendacion.objects.all()
  serializer_class = RecomendacionSerializer
  permission_classes = [IsAuthenticated]
  
  def get_queryset(self):
    queryset = Recomendacion.objects.all()
    evaluacion = self.request.query_params.get('evaluacion')
    if evaluacion is not None:
      queryset = queryset.filter(evaluacion=evaluacion)
    return queryset

  def get_queryset(self):
    queryset = Recomendacion.objects.all()

    # Obtener los parámetros de consulta
    usuario = self.request.query_params.get('usuario')

    # Filtrar por usuario si se proporciona
    if usuario is not None:
      queryset = queryset.filter(usuario=usuario)

    return queryset

class HistorialViewSet(viewsets.ModelViewSet):
  queryset = Historial.objects.all()
  serializer_class = HistorialSerializer
  permission_classes = [IsAuthenticated]

  def get_queryset(self):
    queryset = Historial.objects.all()

    # Obtener los parámetros de consulta
    usuario = self.request.query_params.get('usuario')

    # Filtrar por usuario si se proporciona
    if usuario is not None:
      queryset = queryset.filter(usuario=usuario)

    return queryset

class NotificacionViewSet(viewsets.ModelViewSet):
  queryset = Notificacion.objects.all()
  serializer_class = NotificacionSerializer
  permission_classes = [IsAuthenticated]

  def get_queryset(self):
    queryset = Notificacion.objects.all()
    # Obtener los parámetros de consulta
    usuario = self.request.query_params.get('usuario')
    # Filtrar por usuario si se proporciona
    if usuario is not None:
      queryset = queryset.filter(usuario=usuario)
    return queryset

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def GetUserInfo(request):
  email = request.user
  serializer = UsuarioSerializer(email)
  return Response(serializer.data)

@api_view(['POST'])
@permission_classes([AllowAny])
def SendOtpCode(request):
  email = request.data.get("email")
  if not email:
    return Response({"error": "Email es requerido"}, status=400)

  otp_code = generate_otp()

  # Guardar el OTP en la BD
  OTPCode.objects.create(email=email, code=otp_code)

  # Enviar OTP por correo
  subject = "Tu Código OTP"
  message = f"Tu código OTP es: {otp_code}"
  from_email = "shaquille.montero.vergel123@gmail.com"

  try:
    send_mail(subject, message, from_email, [email], fail_silently=False)
    return Response({"message": "OTP enviado con éxito", "email": email})
  except Exception as e:
    return Response({"error": str(e)}, status=500)

@api_view(['POST'])
@permission_classes([AllowAny])
def VerifyOtpCode(request):
  email = request.data.get("email")
  otp_code = request.data.get("otp")

  if not email or not otp_code:
    return Response({"error": "Email y OTP son requeridos"}, status=400)

  # Buscar el código OTP en la BD
  otp_entry = OTPCode.objects.filter(email=email, code=otp_code).first()

  if otp_entry and otp_entry.is_valid():
    otp_entry.delete()  # Elimina el OTP tras validarlo
    return Response({"message": "OTP válido"})
  else:
    return Response({"error": "OTP inválido o expirado"}, status=400)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def DownloadReportPDF(request):
    print("DownloadReportPDF")
    
    NOMBRE = request.POST.get('NOMBRE')
    CEDULA = request.POST.get('CEDULA')
    PERIODO = request.POST.get('PERIODO')
    MATERIALESCARGADOS = request.POST.get('MATERIALESCARGADOS')
    PREGUNTASGENERADAS = request.POST.get('PREGUNTASGENERADAS')
    EVALUACIONESREALIZADAS = request.POST.get('EVALUACIONESREALIZADAS')

    imagenes = {key: value for key, value in request.FILES.items()}

    datos = {
        'NOMBRE': NOMBRE,
        'CEDULA': CEDULA,
        'PERIODO': PERIODO,
        'MATERIALESCARGADOS': MATERIALESCARGADOS,
        'PREGUNTASGENERADAS': PREGUNTASGENERADAS,
        'EVALUACIONESREALIZADAS': EVALUACIONESREALIZADAS
    }

    template_path = "template.docx"

    # Crear un archivo temporal para el DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_docx_filename = temp_docx.name
    temp_docx.close()

    # Cargar el documento
    doc = Document(template_path)

    # Reemplazar palabras clave en el documento
    for paragraph in doc.paragraphs:
        for key, value in datos.items():
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)

    # Insertar imágenes
    for paragraph in doc.paragraphs:
        for key, file in imagenes.items():
            if key in paragraph.text:
                paragraph.clear()
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                        for chunk in file.chunks():
                            temp_img.write(chunk)
                        temp_img_path = temp_img.name

                    with Image.open(temp_img_path) as image:
                        image.verify()

                    section = doc.sections[0]
                    page_width = section.page_width - section.left_margin - section.right_margin

                    run = paragraph.add_run()
                    run.add_picture(temp_img_path, width=page_width)

                except Exception as e:
                    print(f"Error al procesar la imagen {key}: {e}")

    # Guardar el documento DOCX
    doc.save(temp_docx_filename)

    # Crear un archivo temporal para el PDF
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    temp_pdf_filename = temp_pdf.name
    temp_pdf.close()

    # Convertir DOCX a PDF usando comtypes
    def convert_docx_to_pdf(input_path, output_path):
      comtypes.CoInitialize()  # 🔹 Inicializar COM
      try:
          word = comtypes.client.CreateObject("Word.Application")
          word.Visible = False  # Opcional: ocultar la ventana de Word
          doc = word.Documents.Open(input_path)
          doc.SaveAs(output_path, FileFormat=17)  # 17 es el formato PDF
          doc.Close()
          word.Quit()
      except Exception as e:
          print(f"Error al convertir DOCX a PDF: {e}")
      finally:
          comtypes.CoUninitialize()

    convert_docx_to_pdf(temp_docx_filename, temp_pdf_filename)

    # Eliminar el archivo DOCX temporal
    os.remove(temp_docx_filename)

    # Retornar el archivo PDF como respuesta
    response = FileResponse(open(temp_pdf_filename, 'rb'), as_attachment=True, filename="Reporte.pdf")

    # Eliminar el archivo después de que Django lo envíe
    response["Content-Disposition"] = f'attachment; filename="Reporte.pdf"'
    response["X-Delete-File"] = temp_pdf_filename

    return response